#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Build DOCX for the refined edition."""

import re
import shutil
from pathlib import Path

BUILD   = Path(r"C:\Users\ridge\GitHub_Projects\RidgeAi Projects\GOD of sticks and stones\master_manuscript\07_publication_build")
PROJECT = Path(r"C:\Users\ridge\GitHub_Projects\RidgeAi Projects\GOD of sticks and stones")
OUT_DOCX = BUILD / "review_edition_refined.docx"

SUBS = {
    '\u2014': '--', '\u2013': '-', '\u201c': '"', '\u201d': '"',
    '\u2018': "'", '\u2019': "'", '\u2022': '-',
    '\u2705': '[OK]', '\u26a0': '[!]', '\ufe0f': '',
}

def sanitize(text):
    for k, v in SUBS.items():
        text = text.replace(k, v)
    result = []
    for ch in text:
        try:
            ch.encode('latin-1')
            result.append(ch)
        except Exception:
            result.append('?')
    return ''.join(result)

def strip_md(text):
    text = sanitize(text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*',     r'\1', text)
    text = re.sub(r'`([^`]*)`',     r'\1', text)
    text = re.sub(r'^#+\s*',        '',    text)
    return text.strip()

def main():
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

    md_text = (BUILD / "review_edition_refined.md").read_text(encoding='utf-8')
    doc = Document()

    for section in doc.sections:
        section.page_width  = int(8.5 * 914400)
        section.page_height = int(11  * 914400)
        m = int(1.25 * 914400)
        section.top_margin = section.bottom_margin = m
        section.left_margin = section.right_margin = m

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)

    def add_para(text, align=WD_ALIGN_PARAGRAPH.LEFT,
                 bold=False, italic=False, size=None,
                 space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        return p

    for line in md_text.split('\n'):
        s = line.strip()

        if 'page-break-after' in s:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.add_run().add_break(WD_BREAK.PAGE)
        elif re.match(r'^# [^#]', line):
            add_para(strip_md(line[2:]), align=WD_ALIGN_PARAGRAPH.CENTER,
                     bold=True, size=20, space_before=12, space_after=6)
        elif re.match(r'^## [^#]', line):
            add_para(strip_md(line[3:]), align=WD_ALIGN_PARAGRAPH.CENTER,
                     bold=True, size=15, space_before=10, space_after=4)
        elif re.match(r'^### ', line):
            add_para(strip_md(line[4:]), align=WD_ALIGN_PARAGRAPH.CENTER,
                     bold=True, size=12, space_before=8, space_after=2)
        elif s == '---':
            add_para('-' * 50, space_before=4, space_after=4)
        elif s.startswith('|'):
            cells = [c.strip() for c in s.split('|') if c.strip()]
            if not all(c.replace('-', '').strip() == '' for c in cells):
                add_para('    '.join(cells), size=9, space_after=2)
        elif not s:
            add_para('', space_before=0, space_after=3)
        elif s.startswith('**') and s.endswith('**') and len(s) > 4:
            add_para(strip_md(s[2:-2]), bold=True, space_after=4)
        elif (s.startswith('*') and s.endswith('*')
              and not s.startswith('**') and len(s) > 2):
            add_para(strip_md(s[1:-1]), italic=True, space_after=4)
        elif s.startswith('- ') or s.startswith('* '):
            add_para('  ' + strip_md(s[2:]), space_after=2)
        else:
            add_para(strip_md(s), space_after=4)

    doc.save(str(OUT_DOCX))
    shutil.copy(OUT_DOCX, PROJECT / "review_edition_refined.docx")
    print(f"DOCX saved: {OUT_DOCX}")
    print(f"Size: {OUT_DOCX.stat().st_size:,} bytes")

if __name__ == '__main__':
    main()
