#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Build Review Edition — The God of Sticks and Stones
Produces: review_edition_complete.md / .docx / .pdf
"""

import os, re, sys
from pathlib import Path

# ── PATHS ──────────────────────────────────────────────────────────────────────
PROJECT = Path(r"C:\Users\ridge\GitHub_Projects\RidgeAi Projects\GOD of sticks and stones")
CH = PROJECT / "master_manuscript" / "02_chapters"
OUT = PROJECT / "master_manuscript" / "07_publication_build"

# ── CHAPTER TABLE (num, ordinal, title, filename) ──────────────────────────────
CHAPTERS = [
    (1,  "One",       "Dust and Breath",                        "chapter_01_dust_and_breath.md"),
    (2,  "Two",       "The Tree and the Garden",                "chapter_02_the_tree_and_the_garden.md"),
    (3,  "Three",     "Stone and Witness",                      "chapter_03_stone_and_witness.md"),
    (4,  "Four",      "Seed and Field",                         "chapter_04_seed_and_field.md"),
    (5,  "Five",      "Brick and Empire",                       "chapter_05_brick_and_empire.md"),
    (6,  "Six",       "Wells and Depths",                       "chapter_06_wells_and_depths.md"),
    (7,  "Seven",     "Wilderness and Shepherds",               "chapter_07_wilderness_and_shepherds.md"),
    (8,  "Eight",     "Mountains, Caves, and Lower Places",     "chapter_08_mountains_caves_and_lower_places.md"),
    (9,  "Nine",      "Temple and Tabernacle",                  "chapter_09_temple_and_tabernacle.md"),
    (10, "Ten",       "The Carpenter and the Cross",            "chapter_10_the_carpenter_and_the_cross.md"),
    (11, "Eleven",    "Babel After Babel",                      "chapter_11_babel_after_babel.md"),
    (12, "Twelve",    "The Language of Creation",               "chapter_12_the_language_of_creation.md"),
    (13, "Thirteen",  "The Stone the Builders Rejected",        "chapter_13_the_stone_the_builders_rejected.md"),
    (14, "Fourteen",  "The Living Field",                       "chapter_14_the_living_field.md"),
    (15, "Fifteen",   "The Harvest and the Laborers",           "chapter_15_the_harvest_and_the_laborers.md"),
    (16, "Sixteen",   "Rivers, Wells, and Living Water",        "chapter_16_rivers_wells_and_living_water.md"),
    (17, "Seventeen", "The Final City and the Restored Garden", "chapter_17_the_final_city_and_the_restored_garden.md"),
    # Chapter 18 excluded per review edition assembly instructions
]

PARTS = [
    ("PART I \u2014 FOUNDATIONS OF CREATION",           [1, 2, 3]),
    ("PART II \u2014 BABEL AND THE RISE OF SYSTEMS",    [4, 5, 6, 7]),
    ("PART III \u2014 TEMPLE, CHRIST, AND FULFILLMENT", [8, 9, 10, 11, 12]),
    ("PART IV \u2014 RESTORATION AND THE LIVING FIELD", [13, 14, 15, 16, 17]),
]

# ── FRONT MATTER ───────────────────────────────────────────────────────────────
TITLE    = "THE GOD OF STICKS AND STONES"
SUBTITLE = "Scripture, Creation, and the Quiet Language of God"

DEDICATION_LINES = [
    "For all who have grown weary of noise,",
    "and still hope to hear the quiet things again.",
]

EPIGRAPH_VERSE = "\u201cI tell you, if these were silent,\nthe very stones would cry out.\u201d"
EPIGRAPH_REF   = "\u2014 Luke 19:40"

PREFACE_LINES = [
    "There are truths that speak loudly,",
    "and there are truths that wait quietly beneath the surface of things.",
    "",
    "This book was born from the conviction that scripture continually returns us toward the quiet things.",
    "",
    "Toward rivers.",
    "Toward trees.",
    "Toward fields.",
    "Toward wells.",
    "Toward shepherds.",
    "Toward seeds buried beneath the earth.",
    "Toward stones left standing through generations.",
    "",
    "Again and again,",
    "the biblical story moves through ordinary creation.",
    "",
    "Not because creation itself is divine,",
    "but because creation bears witness.",
    "",
    "The scriptures do not merely present doctrines.",
    "They present patterns.",
    "Movements.",
    "Rhythms.",
    "Ways of seeing.",
    "",
    "The world of scripture is deeply grounded.",
    "Its people walk deserts,",
    "dig wells,",
    "plant vineyards,",
    "watch stars,",
    "follow rivers,",
    "rest beneath trees,",
    "and labor in fields.",
    "",
    "Much of modern life moves differently.",
    "",
    "We live surrounded by speed,",
    "abstraction,",
    "noise,",
    "and systems so dense that we often lose sight of the simple structures that once reminded humanity of dependence,",
    "patience,",
    "and communion.",
    "",
    "This book is not an argument against civilization,",
    "knowledge,",
    "or technology.",
    "",
    "It is not a call toward fear,",
    "withdrawal,",
    "or condemnation.",
    "",
    "Its concern is far more ancient and far more personal.",
    "",
    "It asks whether humanity has slowly lost the ability to:",
    "listen,",
    "notice,",
    "and remain rooted.",
    "",
    "Scripture repeatedly reveals a God who meets humanity:",
    "in gardens,",
    "on mountains,",
    "beside wells,",
    "beneath trees,",
    "along rivers,",
    "in wilderness,",
    "and among laborers.",
    "",
    "The patterns are not accidental.",
    "",
    "This manuscript explores those patterns.",
    "",
    "It reflects upon:",
    "Babel and communion,",
    "depth and elevation,",
    "rootedness and fragmentation,",
    "labor and stewardship,",
    "and Christ as the fulfillment toward which all these symbolic movements quietly point.",
    "",
    "At its heart,",
    "this is a contemplative work.",
    "",
    "It is meant to be read slowly.",
    "",
    "Not rushed.",
    "",
    "The themes intentionally echo and return,",
    "much like rivers revisiting stone,",
    "or seasons revisiting fields.",
    "",
    "The goal is not merely information.",
    "",
    "The goal is attentiveness.",
    "",
    "Because some truths are not grasped through force.",
    "",
    "They are recognized gradually.",
    "",
    "Like roots deepening beneath the surface.",
    "",
    "Like wells hidden beneath dry ground.",
    "",
    "Like seeds rising after long burial.",
    "",
    "And perhaps,",
    "like the quiet language of God that has continued speaking through creation all along.",
]

INTRODUCTION_LINES = [
    "There was a time when humanity lived closer to the ground.",
    "",
    "Closer to rivers.",
    "Closer to seasons.",
    "Closer to the movement of stars.",
    "Closer to the patience required for seeds to rise from buried places.",
    "",
    "The biblical world was not abstract.",
    "",
    "It was tactile.",
    "Grounded.",
    "Rooted in soil,",
    "weather,",
    "stone,",
    "water,",
    "and labor.",
    "",
    "People understood:",
    "how wells were found,",
    "how trees survived drought,",
    "how shepherds followed water,",
    "and how fields depended upon rhythms beyond human control.",
    "",
    "These realities shaped not only survival,",
    "but perception.",
    "",
    "Scripture emerges from this world.",
    "",
    "Its language continually moves through:",
    "gardens,",
    "mountains,",
    "rivers,",
    "wilderness,",
    "bread,",
    "seeds,",
    "vines,",
    "stones,",
    "shepherds,",
    "and storms.",
    "",
    "The patterns repeat so often that they begin to form a kind of symbolic language.",
    "",
    "A language not invented by man,",
    "but embedded throughout creation itself.",
    "",
    "Over time,",
    "human civilization expanded further into:",
    "systems,",
    "abstraction,",
    "speed,",
    "accumulation,",
    "and manufactured identity.",
    "",
    "The scriptures describe this movement repeatedly through Babel.",
    "",
    "Babel is more than a tower.",
    "",
    "It is a pattern.",
    "",
    "A recurring attempt to construct identity,",
    "security,",
    "and greatness apart from communion.",
    "",
    "Yet even after Babel,",
    "scripture continually returns humanity toward:",
    "wells,",
    "shepherds,",
    "fields,",
    "wilderness,",
    "rivers,",
    "and humble places.",
    "",
    "Again and again,",
    "God speaks through the overlooked things.",
    "",
    "A bush burning in wilderness.",
    "A stone set beneath a weary head.",
    "A shepherd boy carrying simple stones.",
    "Water flowing from rock.",
    "A carpenter working with rough wood.",
    "A seed buried before rising.",
    "",
    "The patterns continue until they converge fully in Christ.",
    "",
    "The cornerstone.",
    "The true vine.",
    "The bread of life.",
    "The living water.",
    "The shepherd.",
    "",
    "This book is an attempt to follow those patterns carefully.",
    "",
    "Not to force hidden meanings into every object,",
    "but to observe how scripture repeatedly teaches through creation,",
    "through humility,",
    "and through ordinary realities that modern life often overlooks.",
    "",
    "This is not a rejection of knowledge.",
    "",
    "Knowledge has value.",
    "Craftsmanship has value.",
    "Civilization has value.",
    "",
    "But knowledge is not life.",
    "",
    "And humanity often mistakes accumulation for communion.",
    "",
    "The biblical story continually calls humanity back toward:",
    "dependence,",
    "stewardship,",
    "humility,",
    "and restored relationship.",
    "",
    "The movement is rarely upward first.",
    "",
    "Seeds descend before they rise.",
    "Wells are dug downward before water emerges.",
    "Roots deepen before branches expand.",
    "Humanity falls low before God.",
    "Christ descends before resurrection.",
    "",
    "The pattern repeats across scripture,",
    "creation,",
    "and human experience.",
    "",
    "The modern world often moves too quickly to notice such things.",
    "",
    "Yet the rivers still flow.",
    "The fields still grow.",
    "The stars still testify.",
    "The stones still remain.",
    "",
    "And perhaps the quiet language of God has not disappeared at all.",
    "",
    "Perhaps humanity has simply forgotten how to listen.",
]

NOTE_LINES = [
    "This manuscript is intentionally contemplative in rhythm and structure.",
    "",
    "Certain themes,",
    "images,",
    "and symbols return repeatedly throughout the work.",
    "",
    "These recurrences are intentional.",
    "",
    "The manuscript is designed less like a linear argument,",
    "and more like a slow walk through:",
    "fields,",
    "rivers,",
    "wilderness,",
    "cities,",
    "and finally back toward the garden.",
    "",
    "Readers may benefit from approaching the work slowly,",
    "allowing the symbolic patterns and scriptural echoes to unfold gradually over time.",
]

# ── CLEAN CHAPTER LINES ────────────────────────────────────────────────────────
def clean_chapter_lines(num, ordinal, title, filename):
    """Read a chapter file, strip redundant title/continuation headers, return clean lines."""
    path = CH / filename
    raw = path.read_text(encoding="utf-8").splitlines()

    # Remove trailing trailing blank lines
    while raw and not raw[-1].strip():
        raw.pop()

    cleaned = []
    i = 0
    while i < len(raw):
        line = raw[i]

        # Strip: "# THE GOD OF STICKS AND STONES"
        if line.strip() == "# THE GOD OF STICKS AND STONES":
            i += 1
            continue

        # Strip: "## Chapter Two — The Tree and the Garden (Continuation)"
        if "(Continuation)" in line:
            i += 1
            continue

        # Fix: Ch02 Part-1 opens with "# Chapter Two…" (single #) — normalize to ##
        if re.match(r'^# Chapter ', line) and not line.startswith('##'):
            normalized = "## " + line[2:].strip()
            cleaned.append(normalized)
            i += 1
            continue

        # Normalize any ## heading that has the book title baked in
        if line.startswith("## ") and "THE GOD OF STICKS" in line:
            i += 1
            continue

        cleaned.append(line)
        i += 1

    # Ensure first non-blank line is the chapter heading
    first_content = next((l for l in cleaned if l.strip()), "")
    if not first_content.startswith("## "):
        cleaned.insert(0, f"## Chapter {ordinal} \u2014 {title}")
        cleaned.insert(1, "")

    return cleaned


# ══════════════════════════════════════════════════════════════════════════════
# 1.  BUILD MARKDOWN
# ══════════════════════════════════════════════════════════════════════════════
def build_markdown():
    md = []

    def pb():
        """Page-break marker visible to converters; ignored by plain MD viewers."""
        md.append("")
        md.append('<div style="page-break-after: always;"></div>')
        md.append("")

    def rule():
        md.append("")
        md.append("---")
        md.append("")

    # ── Title Page ────────────────────────────────────────────────────────────
    md += [
        f"# {TITLE}",
        "",
        f"## {SUBTITLE}",
        "",
        "",
        "*Version 1 \u2014 Review Edition*",
        "",
    ]
    pb()

    # ── Dedication ────────────────────────────────────────────────────────────
    md += ["### Dedication", ""]
    md += DEDICATION_LINES
    md += [""]
    pb()

    # ── Epigraph ─────────────────────────────────────────────────────────────
    md += ["### Epigraph", ""]
    for ln in EPIGRAPH_VERSE.split("\n"):
        md.append(f"*{ln}*")
    md += ["", f"*{EPIGRAPH_REF}*", ""]
    pb()

    # ── Preface ───────────────────────────────────────────────────────────────
    md += ["## Preface", ""]
    md += PREFACE_LINES
    md += [""]
    pb()

    # ── Introduction ──────────────────────────────────────────────────────────
    md += ["## Introduction", ""]
    md += INTRODUCTION_LINES
    md += [""]
    pb()

    # ── A Note on Reading ─────────────────────────────────────────────────────
    md += ["### A Note on Reading This Book", ""]
    md += NOTE_LINES
    md += [""]
    pb()

    # ── Table of Contents ─────────────────────────────────────────────────────
    md += ["## Table of Contents", ""]
    md += [
        "**Front Matter**",
        "",
        "Dedication",
        "Epigraph",
        "Preface",
        "Introduction",
        "A Note on Reading This Book",
        "",
    ]
    for part_title, ch_nums in PARTS:
        md.append(f"**{part_title}**")
        md.append("")
        for n, ord_, t, _ in CHAPTERS:
            if n in ch_nums:
                md.append(f"Chapter {ord_} \u2014 {t}")
        md.append("")

    md += [
        "**Back Matter**",
        "",
        "Appendix I \u2014 Recurring Biblical Symbols",
        "Appendix II \u2014 Scripture Pathways",
        "Appendix III \u2014 Babel and Communion Study Framework",
        "Appendix IV \u2014 The Language of Creation",
        "Symbol Index",
        "Scripture Index",
        "Author Intent Statement",
        "Closing Reflection *(to be written)*",
        "Acknowledgments *(to be written)*",
        "",
    ]
    pb()

    # ── MAIN BODY ─────────────────────────────────────────────────────────────
    for part_title, ch_nums in PARTS:
        md += [f"# {part_title}", ""]
        pb()

        for num, ordinal, title, fname in CHAPTERS:
            if num not in ch_nums:
                continue
            lines = clean_chapter_lines(num, ordinal, title, fname)
            md += lines
            md += [""]
            rule()
            pb()

    # ── BACK MATTER ───────────────────────────────────────────────────────────
    # -- Appendix I
    md += [
        "# Back Matter",
        "",
        "## Appendix I \u2014 Recurring Biblical Symbols",
        "",
        "### Trees",
        "",
        "**Core themes:** rootedness, covenant, life, restoration, cross imagery",
        "",
        "**Key scriptures:** Genesis 2; Psalm 1; Isaiah 61; John 15; Revelation 22",
        "",
        "### Stone",
        "",
        "**Core themes:** witness, permanence, remembrance, foundation, cornerstone",
        "",
        "**Key scriptures:** Genesis 28; 1 Samuel 17; Psalm 118:22; Luke 19:40; 1 Peter 2",
        "",
        "### Wells and Water",
        "",
        "**Core themes:** hidden life, dependence, descent, communion, living water, thirst",
        "",
        "**Key scriptures:** Genesis 21, 26, 29; Exodus 17; John 4; John 7:37; Revelation 22",
        "",
        "### Seeds and Fields",
        "",
        "**Core themes:** burial before life, hiddenness, stewardship, patience, harvest",
        "",
        "**Key scriptures:** Genesis 8; Matthew 13; John 12:24; Galatians 6:7\u20139",
        "",
        "### Wilderness",
        "",
        "**Core themes:** dependence, purification, encounter, formation",
        "",
        "**Key scriptures:** Exodus 16\u201317; 1 Kings 19; Matthew 4",
        "",
        "### Mountains and Caves",
        "",
        "**Core themes:** humility before God, encounter, revelation, hiddenness",
        "",
        "**Key scriptures:** Exodus 19\u201320; 1 Kings 19; Matthew 5:1; Mark 9:2",
        "",
    ]
    pb()

    # -- Appendix II
    md += [
        "## Appendix II \u2014 Scripture Pathways",
        "",
        "**Pathway 1 \u2014 The Garden to the Restored Garden**",
        "Genesis 1\u20132 \u2192 Isaiah 35, 55 \u2192 John 20 \u2192 Revelation 21\u201322",
        "",
        "**Pathway 2 \u2014 Wells, Rivers, and Living Water**",
        "Genesis 21, 26, 29 \u2192 Exodus 17 \u2192 Psalm 46 \u2192 Isaiah 43 \u2192 John 4 \u2192 John 7 \u2192 Revelation 22",
        "",
        "**Pathway 3 \u2014 The Stone Motif**",
        "Genesis 28 \u2192 1 Samuel 17 \u2192 Psalm 118 \u2192 Matthew 21:42 \u2192 1 Peter 2 \u2192 Revelation 21",
        "",
        "**Pathway 4 \u2014 Shepherds and Stewardship**",
        "Genesis 46\u201347 \u2192 Exodus 3 \u2192 1 Samuel 16\u201317 \u2192 Psalm 23 \u2192 John 10",
        "",
        "**Pathway 5 \u2014 Babel and the City**",
        "Genesis 11 \u2192 Isaiah 13\u201314 \u2192 Daniel \u2192 Revelation 17\u201318 \u2192 Revelation 21",
        "",
        "**Pathway 6 \u2014 Bread, Seed, and Harvest**",
        "Genesis 8:22 \u2192 Leviticus 23 \u2192 Ruth \u2192 Matthew 6:11 \u2192 John 6 \u2192 Revelation 14",
        "",
    ]
    pb()

    # -- Appendix III
    md += [
        "## Appendix III \u2014 Babel and Communion Study Framework",
        "",
        "**What This Manuscript Does Not Argue**",
        "",
        "That cities are evil.",
        "That knowledge is evil.",
        "That technology is inherently corrupt.",
        "That civilization should be abandoned.",
        "",
        "**What This Manuscript Does Explore**",
        "",
        "The recurring orientation of human systems toward self-exaltation apart from communion.",
        "The contrast between structures that support dependence on God versus systems that attempt to replace it.",
        "The spiritual consequences of disconnection from creation\u2019s witness.",
        "The restoration toward which Christ points \u2014 which includes a renewed city.",
        "",
        "**The Central Tension**",
        "",
        "Communion-oriented structures (garden, well, altar, field, shepherd)",
        "versus self-exalting systems detached from dependence (tower, empire, brick, manufactured permanence).",
        "",
        "**Important Balance**",
        "",
        "Scripture ends with a restored city (Revelation 21\u201322).",
        "The issue was never civilization itself.",
        "The issue was separation from communion.",
        "The kingdom\u2019s goal is restoration, not destruction.",
        "",
    ]
    pb()

    # -- Appendix IV
    md += [
        "## Appendix IV \u2014 The Language of Creation",
        "",
        "**Seeds and Burial** \u2014 Life entering hiddenness before emergence; resurrection pattern in created order.",
        "",
        "**Rivers and Descent** \u2014 Water flowing downward; life found through humility.",
        "",
        "**Trees and Rootedness** \u2014 Hidden foundations sustaining visible life; dependence before fruitfulness.",
        "",
        "**Gravity and Humility** \u2014 Created order\u2019s witness to limitation; towers returning to dust.",
        "",
        "**Wells and Hidden Life** \u2014 Descent required before life is found; encounter at lowered places.",
        "",
        "**Bread and Dependence** \u2014 Daily dependence; the body\u2019s need; Christ as bread of life.",
        "",
        "**Stars and Order** \u2014 Scale and wonder; the vastness of creation humbling humanity.",
        "",
        "**Shepherding and Stewardship** \u2014 Relational leadership; care for the individual within the flock.",
        "",
        "*Note: Creation is not worshiped. Creation functions as witness.*",
        "*The manuscript continually points through creation, not away from God toward creation itself.*",
        "",
    ]
    pb()

    # -- Symbol Index
    md += [
        "## Symbol Index",
        "",
        "| Symbol | Chapters |",
        "|--------|---------|",
        "| Babel / Tower | 1, 5, 11, 13, 17 |",
        "| Bread | 9, 10, 14, 15 |",
        "| Caves | 8 |",
        "| Cross / Wood | 2, 10 |",
        "| Dust | 1, 3, 8 |",
        "| Field | 1, 4, 14, 15 |",
        "| Garden | 1, 2, 9, 10, 17 |",
        "| Gravity / Loweredness | 1, 8, 10, 11, 13 |",
        "| Mountains | 8 |",
        "| River | 6, 9, 16, 17 |",
        "| Seeds | 1, 4, 10, 12, 14, 15 |",
        "| Shepherds | 7, 9, 15 |",
        "| Stone / Cornerstone | 3, 8, 10, 13 |",
        "| Temple | 9 |",
        "| Trees | 1, 2, 9, 10 |",
        "| Wilderness | 7, 8 |",
        "| Wells | 2, 5, 6, 16 |",
        "",
    ]
    pb()

    # -- Scripture Index
    md += [
        "## Scripture Index",
        "",
        "| Scripture | Theme | Chapter(s) |",
        "|-----------|-------|-----------|",
        "| Genesis 1\u20132 | Garden, creation, trees, rivers | 1, 2, 16 |",
        "| Genesis 11 | Babel, brick, tower | 1, 5, 11 |",
        "| Genesis 28 | Jacob, stone witness, encounter | 3, 13 |",
        "| Exodus 17 | Water from rock, wilderness | 3, 6, 7 |",
        "| 1 Samuel 17 | David, stone, empire confronted | 13 |",
        "| Psalm 1 | Tree planted by water | 2, 14 |",
        "| Psalm 19 | Heavens declaring glory | 12 |",
        "| Psalm 118:22 | Rejected cornerstone | 13 |",
        "| Isaiah 55 | Water for the thirsty | 16 |",
        "| John 4 | Living water, Samaritan woman | 6, 16 |",
        "| John 10 | Good Shepherd | 7, 15 |",
        "| John 12:24 | Grain of wheat | 4, 10 |",
        "| John 15 | True vine | 2, 12 |",
        "| Luke 19:40 | Stones crying out | 3, 10 |",
        "| Matthew 13 | Parables of the kingdom | 4, 14, 15 |",
        "| 1 Peter 2:4\u20135 | Living stones | 3, 13 |",
        "| Revelation 21\u201322 | New Jerusalem, river, tree of life | 9, 17 |",
        "",
    ]
    pb()

    # -- Author Intent Statement (from existing file)
    intent_file = OUT / "author_intent_statement.md"
    if intent_file.exists():
        intent_text = intent_file.read_text(encoding="utf-8")
        # Strip the YAML-ish header lines before the actual prose
        intent_lines = intent_text.splitlines()
        # Remove header/title lines at top (lines starting with #)
        start = 0
        for idx, l in enumerate(intent_lines):
            if l.startswith("## STATEMENT OF INTENT"):
                start = idx
                break
        md += ["## Author Intent Statement", ""]
        md += intent_lines[start:]
        md += [""]
    pb()

    # -- Closing Reflection placeholder
    md += [
        "## Closing Reflection",
        "",
        "*(To be written.)*",
        "",
        "The field remains unfinished.",
        "",
        "Rivers continue flowing.",
        "",
        "The invitation remains open.",
        "",
        "Communion is still available.",
        "",
        "The reader now returns to ordinary life differently.",
        "",
    ]
    pb()

    # -- Acknowledgments placeholder
    md += [
        "## Acknowledgments",
        "",
        "*(To be written by the author.)*",
        "",
        "Tone: sincere, humble, grounded, restrained.",
        "",
        "Suggested themes: gratitude toward faithful laborers, quiet mentors, family, scripture, and generations who preserved wisdom.",
        "",
    ]

    # ── EDITORIAL NOTICE ──────────────────────────────────────────────────────
    md += [
        "",
        "---",
        "",
        "## Editorial Notice",
        "",
        "**Version 1 \u2014 Review Edition**",
        "",
        "This document is assembled for author readthrough, beta reader review,",
        "pacing evaluation, theological review, and editorial annotation.",
        "",
        "**Note:** Chapter 18 (\u201cThe God of Sticks and Stones\u201d) is not included in this",
        "review edition per assembly instructions. It exists as a complete chapter",
        "in `master_manuscript/02_chapters/chapter_18_the_god_of_sticks_and_stones.md`",
        "and should be reviewed for inclusion as the manuscript\u2019s closing chapter.",
        "",
        "**Note:** Closing Reflection and Acknowledgments are placeholder sections.",
        "Full content to be written by the author.",
        "",
    ]

    return "\n".join(md)


# ══════════════════════════════════════════════════════════════════════════════
# 2.  BUILD DOCX
# ══════════════════════════════════════════════════════════════════════════════
def build_docx(md_text):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # --- Page setup: letter, 1-inch margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # --- Adjust Normal style: Times New Roman 11pt, 1.5 line spacing
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)

    def set_line_spacing(para, spacing_val=276):  # 276 twips = 1.5 lines (184 = single)
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), str(spacing_val))
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)

    def add_page_break(doc):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(docx_page_break())
        return p

    # Helper for page break
    from docx.oxml import OxmlElement
    def docx_page_break():
        from docx.oxml.ns import qn
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        return br

    def add_pb(doc):
        p = doc.add_paragraph()
        run = p.add_run()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)

    def para(doc, text, style='Normal', bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, size=None, space_before=0, space_after=3):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            if size:
                run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
        set_line_spacing(p)
        return p

    def heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER):
        sizes = {1: 22, 2: 16, 3: 13}
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 12))
        run.font.name = 'Times New Roman'
        return p

    def small_space(doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run('')
        run.font.size = Pt(6)

    # ── Parse markdown and emit DOCX ─────────────────────────────────────────
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Page break marker
        if '<div style="page-break-after: always;"></div>' in line:
            add_pb(doc)
            i += 1
            continue

        # H1 heading
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            heading(doc, text, level=1)
            i += 1
            continue

        # H2 heading
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            heading(doc, text, level=2)
            i += 1
            continue

        # H3 heading
        if line.startswith('### '):
            text = line[4:].strip()
            heading(doc, text, level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Table row (in indexes)
        if line.startswith('|'):
            # Simple table: render each cell as a tab-separated line
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells and not all(c.startswith('-') for c in cells):
                txt = '    '.join(cells)
                p = para(doc, txt, size=9, space_after=1)
            i += 1
            continue

        # Empty line
        if not line.strip():
            small_space(doc)
            i += 1
            continue

        # Italic text (*...*) — used for scripture quotes and notes
        if line.strip().startswith('*') and line.strip().endswith('*'):
            text = line.strip()[1:-1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(text)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            set_line_spacing(p)
            i += 1
            continue

        # Bold text (**...**)
        if line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            p = para(doc, text, bold=True, space_before=6, space_after=2)
            i += 1
            continue

        # Regular body text
        text = line.strip()
        # Strip inline markdown bold/italic for plain output
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*',     r'\1', text)
        p = para(doc, text, space_after=3)
        i += 1

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# 3.  BUILD PDF  (fpdf2 — Times core font with Unicode substitution)
# ══════════════════════════════════════════════════════════════════════════════
def build_pdf(md_text):
    from fpdf import FPDF

    # Substitute characters that Latin-1 core fonts can't render
    SUBS = {
        '\u2014': '--',   # em dash
        '\u2013': '-',    # en dash
        '\u201c': '"',    # left double quote
        '\u201d': '"',    # right double quote
        '\u2018': "'",    # left single quote
        '\u2019': "'",    # right single quote
        '\u2192': '->',   # arrow
        '\u2019': "'",
        '\u00e9': 'e',    # e accent
    }

    def clean(text):
        for k, v in SUBS.items():
            text = text.replace(k, v)
        return text

    class ManuscriptPDF(FPDF):
        def header(self):
            pass  # No running header for review edition

        def footer(self):
            self.set_y(-15)
            self.set_font('Times', 'I', 9)
            self.cell(0, 10, f'Page {self.page_no()}', align='C')

    pdf = ManuscriptPDF(orientation='P', unit='mm', format='Letter')
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(left=25.4, top=25.4, right=25.4)
    pdf.add_page()

    def h1(text):
        pdf.ln(6)
        pdf.set_font('Times', 'B', 20)
        pdf.multi_cell(0, 10, clean(text), align='C')
        pdf.ln(4)

    def h2(text):
        pdf.ln(5)
        pdf.set_font('Times', 'B', 15)
        pdf.multi_cell(0, 8, clean(text), align='C')
        pdf.ln(3)

    def h3(text):
        pdf.ln(4)
        pdf.set_font('Times', 'B', 12)
        pdf.multi_cell(0, 7, clean(text), align='C')
        pdf.ln(2)

    def body(text, italic=False):
        style = 'I' if italic else ''
        pdf.set_font('Times', style, 11)
        pdf.multi_cell(0, 6, clean(text), align='L')

    def blank():
        pdf.ln(4)

    def rule_line():
        pdf.ln(3)
        pdf.set_draw_color(150, 150, 150)
        x = pdf.get_x()
        y = pdf.get_y()
        pdf.line(25.4, y, 186, y)
        pdf.ln(4)

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if '<div style="page-break-after: always;"></div>' in line:
            pdf.add_page()
            i += 1
            continue

        if line.startswith('# ') and not line.startswith('## '):
            h1(line[2:].strip())
            i += 1
            continue

        if line.startswith('## ') and not line.startswith('### '):
            h2(line[3:].strip())
            i += 1
            continue

        if line.startswith('### '):
            h3(line[4:].strip())
            i += 1
            continue

        if line.strip() == '---':
            rule_line()
            i += 1
            continue

        if line.startswith('|'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells and not all(c.replace('-','').replace(' ','') == '' for c in cells):
                txt = '    '.join(cells)
                pdf.set_font('Times', '', 9)
                pdf.multi_cell(0, 5, clean(txt), align='L')
            i += 1
            continue

        if not line.strip():
            blank()
            i += 1
            continue

        text = line.strip()
        is_italic = text.startswith('*') and text.endswith('*') and len(text) > 2
        if is_italic:
            text = text[1:-1]
        is_bold = text.startswith('**') and text.endswith('**') and len(text) > 4
        if is_bold:
            text = text[2:-2]
            pdf.set_font('Times', 'B', 11)
            pdf.multi_cell(0, 6, clean(text), align='L')
            i += 1
            continue

        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*',     r'\1', text)
        body(text, italic=is_italic)
        i += 1

    return pdf


# ══════════════════════════════════════════════════════════════════════════════
# 4.  STATS
# ══════════════════════════════════════════════════════════════════════════════
def compute_stats(md_text):
    # Strip markdown syntax for word count
    plain = re.sub(r'<[^>]+>', '', md_text)
    plain = re.sub(r'[#*|`\-]', ' ', plain)
    words = len(plain.split())
    # Chapter word counts
    chapter_words = {}
    for num, ordinal, title, fname in CHAPTERS:
        path = CH / fname
        raw = path.read_text(encoding='utf-8')
        raw = re.sub(r'[#*|`\-]', ' ', raw)
        chapter_words[num] = (title, len(raw.split()))

    total_chapter_words = sum(v[1] for v in chapter_words.values())
    return words, chapter_words, total_chapter_words


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    OUT.mkdir(parents=True, exist_ok=True)

    print("Building markdown...")
    md_text = build_markdown()
    md_path = OUT / "review_edition_complete.md"
    md_path.write_text(md_text, encoding='utf-8')
    print(f"  -> {md_path}")

    print("Building DOCX...")
    doc = build_docx(md_text)
    docx_path = OUT / "review_edition_complete.docx"
    doc.save(str(docx_path))
    print(f"  -> {docx_path}")

    print("Building PDF...")
    try:
        pdf = build_pdf(md_text)
        pdf_path = OUT / "review_edition_complete.pdf"
        pdf.output(str(pdf_path))
        print(f"  -> {pdf_path}")
        pdf_ok = True
    except Exception as e:
        print(f"  PDF error: {e}")
        pdf_ok = False
        pdf_path = None

    # Also copy to root
    import shutil
    shutil.copy(md_path,   PROJECT / "review_edition_complete.md")
    shutil.copy(docx_path, PROJECT / "review_edition_complete.docx")
    if pdf_ok:
        shutil.copy(pdf_path, PROJECT / "review_edition_complete.pdf")

    # ── Stats & Summary ────────────────────────────────────────────────────────
    print("Computing stats...")
    total_words, ch_words, body_words = compute_stats(md_text)

    # Estimate page counts
    # Review edition (12pt Times, 1.5 spacing, 1.25" margins, letter)
    # Effective ~220 words/page for dense prose, ~120 for highly spaced
    # This manuscript is highly line-broken, estimate ~150 words/page for body
    estimated_pages_body = round(body_words / 150)
    estimated_pages_front = 12
    estimated_pages_back  = 15
    estimated_pages_total = estimated_pages_body + estimated_pages_front + estimated_pages_back

    summary_lines = [
        "# Review Edition Summary",
        "## The God of Sticks and Stones",
        "### Version 1 — Review Edition",
        "",
        f"**Assembly Date:** 2026-05-23",
        "",
        "---",
        "",
        "## TOTAL WORD COUNT",
        "",
        f"| Component | Words |",
        f"|-----------|-------|",
        f"| Main body (Chapters 1–17) | {body_words:,} |",
        f"| Full document (with front/back matter) | {total_words:,} |",
        "",
        "---",
        "",
        "## ESTIMATED PAGE COUNT",
        "",
        f"| Section | Estimated Pages |",
        f"|---------|----------------|",
        f"| Front matter (title, dedication, epigraph, preface, intro, note, TOC) | ~{estimated_pages_front} |",
        f"| Main body (17 chapters, 4 parts) | ~{estimated_pages_body} |",
        f"| Back matter (appendices, indexes, intent statement) | ~{estimated_pages_back} |",
        f"| **Total estimated pages** | **~{estimated_pages_total}** |",
        "",
        "*Estimate based on Times New Roman 11pt, 1.5 line spacing, 1.25\" margins, Letter format.*",
        "*This manuscript's contemplative short-line format yields lower word density per page than standard prose.*",
        "",
        "---",
        "",
        "## CHAPTER INVENTORY",
        "",
        "| # | Chapter Title | Word Count | Status |",
        "|---|--------------|-----------|--------|",
    ]

    for num, ordinal, title, _ in CHAPTERS:
        t, wc = ch_words[num]
        summary_lines.append(f"| {num} | {title} | {wc:,} | ✅ Complete |")

    summary_lines += [
        f"| 18 | The God of Sticks and Stones | *(excluded from review edition)* | ⚠️ Exists — not included |",
        "",
        f"**Total chapters assembled:** 17",
        f"**Chapter 18 status:** Complete file exists at `master_manuscript/02_chapters/chapter_18_the_god_of_sticks_and_stones.md`",
        "",
        "---",
        "",
        "## FILES GENERATED",
        "",
        f"| File | Location | Status |",
        f"|------|----------|--------|",
        f"| review_edition_complete.md | Project root + 07_publication_build/ | ✅ Created |",
        f"| review_edition_complete.docx | Project root + 07_publication_build/ | ✅ Created |",
        f"| review_edition_complete.pdf | Project root + 07_publication_build/ | {'✅ Created' if pdf_ok else '⚠️ Generation error — see note'} |",
        "",
        "---",
        "",
        "## UNRESOLVED ISSUES",
        "",
        "| # | Issue | Location | Severity |",
        "|---|-------|----------|----------|",
        "| 1 | Chapter 18 excluded from review edition | Assembly instructions | NOTE |",
        "| 2 | Closing Reflection not yet written | Back matter | MODERATE |",
        "| 3 | Acknowledgments not yet written | Back matter | LOW |",
        "| 4 | Introduction is draft length (~700 words); target 3,000–6,000 | Front matter | LOW |",
        "| 5 | Dedication: author must select from 5 options | Front matter | LOW |",
        "| 6 | Epigraph: author must confirm Luke 19:40 selection | Front matter | LOW |",
        "| 7 | Self-referential passage in Chapter 15 (~line 99–110) | Ch. 15 body | AUTHOR DECISION |",
        "| 8 | Appendices I–IV are frameworks only; full prose not written | Back matter | MODERATE |",
        "| 9 | Chapter 2 file seam (merged from 2 source files) | Ch. 2 | MINOR |",
        "| 10 | 17 vs 18 chapter structure decision | Overall | AUTHOR DECISION |",
        "",
        "---",
        "",
        "## SECTIONS NEEDING FUTURE REFINEMENT",
        "",
        "| Priority | Section | Issue |",
        "|----------|---------|-------|",
        "| 1 | Chapter 15 | Self-referential 'the manuscript has continually contrasted' passage — decide to preserve or revise |",
        "| 1 | Chapter 18 | Whether to include as final chapter of the main body |",
        "| 2 | Ch. 2 → Ch. 3 transition | Slightly abrupt shift from tree imagery to stone; consider bridging sentence |",
        "| 2 | Ch. 4 → Ch. 5 transition | Field → Brick/Empire shift; brief connecting image recommended |",
        "| 2 | Ch. 12 → Ch. 13 transition | Stars/cosmic ending into 'Humanity is always building' — slight tonal shift |",
        "| 3 | Modern civilization critique phrasing | 'Speed, noise, screens, artificial light' catalog appears 6+ times nearly verbatim across chapters |",
        "| 3 | Well/Tower contrast language | Chs. 16–18 use very similar language; audit for freshness |",
        "| 3 | Introduction expansion | Current draft (~700 words) may benefit from expansion |",
        "| 4 | Chapter 5 (Brick and Empire) | Slightly more critical in tone than surrounding chapters; author may wish to review |",
        "| 4 | Chapter 11 (Babel After Babel) | Highest density of modern civilization critique; ensure it advances beyond Ch. 5 |",
        "",
        "---",
        "",
        "## MISSING CONTENT",
        "",
        "| Item | Notes |",
        "|------|-------|",
        "| Closing Reflection (full prose) | Framework: 'The field remains unfinished. Rivers continue flowing.' Needs authorial voice. |",
        "| Acknowledgments (full prose) | Framework exists. Author to write. |",
        "| Appendices I–IV (full prose) | All four appendix outlines complete. ~500 words each needed. |",
        "| Full Symbol Index | Sample entries assembled. Complete compilation from all 18 chapters needed. |",
        "| Full Scripture Index | Sample entries assembled. Complete compilation needed. |",
        "| Dedicated 'Seed and the Fall' chapter | Original plan target Ch. 2; never written as standalone. Fall is addressed in Chs. 1, 2, 14. Author must decide. |",
        "",
        "---",
        "",
        "## FORMATTING CONCERNS",
        "",
        "| Concern | Notes |",
        "|---------|-------|",
        "| PDF Unicode | PDF uses ASCII substitutions for em dashes and smart quotes (review edition only). Final typeset requires proper font embedding. |",
        "| DOCX table formatting | Index tables render as tab-separated text rows in DOCX review edition. Will require proper table formatting for final publication. |",
        "| Chapter 2 seam | The merged Part 1 + Continuation flows well. Review the transition between the two parts during readthrough (approximately 60% through Chapter 2). |",
        "| Short-line spacing | The manuscript's contemplative short-line format results in significantly more pages than equivalent word count in standard prose. This is intentional and correct for the genre. |",
        "",
        "---",
        "",
        "## THEOLOGICAL AND VOICE INTEGRITY",
        "",
        "No theology was rewritten in this assembly.",
        "No prose was modernized.",
        "No contemplative pacing was compressed.",
        "No symbolic structure was altered.",
        "",
        "All chapter content is reproduced exactly from source files,",
        "with the following technical normalizations only:",
        "- Removed repeated book title headers from chapter openings",
        "- Removed 'Chapter Two — The Tree and the Garden (Continuation)' sub-header",
        "- Normalized one single-# heading to double-# heading style",
        "",
        "---",
        "",
        "## OVERALL ASSESSMENT",
        "",
        "The Version 1 Review Edition is assembled and ready for human readthrough.",
        "",
        "The manuscript body is in excellent condition.",
        "The literary voice is consistent and distinctive throughout.",
        "The symbolic architecture is intact and coherent.",
        "The Christ-centered convergence is present in all 17 assembled chapters.",
        "",
        "Recommended reading order for review:",
        "1. Read the Preface and Introduction first for orientation",
        "2. Read Part I (Chs. 1–3) as a single sitting to establish the symbolic register",
        "3. Read Part II (Chs. 4–7) noting the Babel/communion tension building",
        "4. Read Part III (Chs. 8–12) as the theological convergence",
        "5. Read Part IV (Chs. 13–17) as the restoration arc",
        "6. Note: Ch. 18 ('The God of Sticks and Stones') exists separately and",
        "   should be read as the culminating chapter",
        "",
        "---",
        "",
        "*Review edition assembled 2026-05-23*",
    ]

    summary_text = "\n".join(summary_lines)
    summary_path = OUT / "review_edition_summary.md"
    summary_path.write_text(summary_text, encoding='utf-8')
    shutil.copy(summary_path, PROJECT / "review_edition_summary.md")
    print(f"  -> {summary_path}")

    print()
    print("=" * 60)
    print("REVIEW EDITION ASSEMBLY COMPLETE")
    print("=" * 60)
    print(f"  Markdown:  review_edition_complete.md")
    print(f"  DOCX:      review_edition_complete.docx")
    print(f"  PDF:       {'review_edition_complete.pdf' if pdf_ok else 'FAILED (see error above)'}")
    print(f"  Summary:   review_edition_summary.md")
    print(f"  Body words: {body_words:,}")
    print(f"  Total words: {total_words:,}")
    print(f"  Est. pages: ~{estimated_pages_total}")
    print("=" * 60)
