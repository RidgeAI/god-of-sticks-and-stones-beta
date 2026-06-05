"""
Review Packet Generator
The God of Sticks and Stones — Chapter 18 and Closing Reflections
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ────────────────────────────────────────────────────────────────────

PROJECT_ROOT = Path(r"C:\Users\ridge\GitHub_Projects\RidgeAi Projects\GOD of sticks and stones")
OUTPUT_DIR   = PROJECT_ROOT / "master_manuscript" / "07_publication_build" / "review_packets"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ─── Content ──────────────────────────────────────────────────────────────────

BOOK_TITLE    = "The God of Sticks and Stones"
BOOK_SUBTITLE = "Scripture, Creation, and the Quiet Language of God"

CHAPTER_18_HEADING = "Chapter Eighteen — The God of Sticks and Stones"

CHAPTER_18 = """\
And throughout scripture,
God continually speaks through ordinary things.

Stick.
Stone.
Seed.
Field.
Water.
Tree.

Simple things.
Foundational things.
Things humanity often overlooks while searching for greatness elsewhere.

Perhaps modern humanity has not truly become more advanced than previous generations.

Perhaps humanity has simply become more distracted.

The pace increased.
The towers grew taller.
The systems became more complex.

But the deepest human questions remain unchanged.

What is life?
What is truth?
What is worth building?
What foundation can endure?
What truly satisfies thirst?

Scripture answers these questions not first through abstraction,
but through living symbols grounded in creation itself.

Water.
Seed.
Bread.
Stone.
Cross.
River.
Tree.

Again and again,
God speaks through things close enough to touch.

This keeps the kingdom near the ground.

Near enough for children.
Near enough for laborers.
Near enough for the poor.
Near enough for farmers,
shepherds,
and ordinary people.

Empire often speaks the language of complexity.
The kingdom often speaks the language of creation.

The stars and the seed tell the same story.
The well and the cross tell the same story.
The river and the tree tell the same story.

Life comes through communion,
through surrender,
through rootedness,
and through receiving what humanity cannot manufacture for itself.

This is why Christ becomes the cornerstone.

And still,
He remains close enough to kneel beside the ground.

Perhaps this is the final invitation of the kingdom.

Not upward striving toward self-made transcendence.

But downward rootedness into reality,
communion,
and love.

To become once again:
- living soil,
- faithful laborers,
- rooted trees,
- people willing to receive living water,
- builders who remember the cornerstone,
- and stewards humble enough to walk with God through the field.

For beneath every empire,
beyond every tower,
and underneath all the noise of human history,
the older realities still remain.

Stone.
Tree.
Water.
Seed.
Dust.
Breath.

And the God who has always spoken through them.\
"""

CLOSING_A_TITLE = "Closing Reflection A — The Laborer's Return"
CLOSING_A = """\
The work is still waiting for you.

Not the extraordinary kind — just the ordinary. The meetings and the dishes. The conversation that doesn't quite resolve. The commute. The quiet morning. The neighbor whose name you keep meaning to learn.

You have been walking through something slow and old. And now the ordinary day is in front of you again, with its small weight, asking for your presence.

This is not a descent.

Those who worked close to the ground understood that faithfulness rarely meant stepping away from ordinary life. It meant bringing a different quality of attention into it. Showing up. Tending what was given to them to tend.

You don't have to be extraordinary at this.

You only need to remain present.

And the one who worked with his hands, who ate with people who weren't certain they belonged at the table, has always been willing to meet you in the ordinary hours — in the commute, in the kitchen, in the work that will still be there tomorrow morning.

He did not require spectacle from the people he met along the road.

He is not requiring it now.

So go back to the work.

Quietly, and with purpose — the way a laborer returns after a long rest. The task was always ordinary. And you were always enough to carry it.\
"""

CLOSING_B_TITLE = "Closing Reflection B — The Tuning Fork"
CLOSING_B = """\
Something has slowed in you, briefly.

It may not stay in exactly this form. The noise will return — it always does. The pace will pick up. The morning after this one will feel more like all the others.

But something will remain.

Not a system. Not a practice. Just the quiet knowledge that the world was speaking before you found this book, and that it will go on speaking when you set it down.

There is a stone somewhere near where you live. Or the sound of wind through something — leaves against glass, a window screen, the gap beneath a door on a cold morning. Small sounds. Ordinary sounds. The kind you stopped registering a long time ago.

They were always there.

That is what this slow walk has been pointing toward. Not toward something new, but back toward what was already present. The language was not lost. Only your capacity to hear it had grown difficult.

You have been quiet for a little while now.

That is enough.

The speaking continues — whether you are attending or not, whether the day is easy or not, whether you remember this or not. It was never waiting for the right kind of person to hear it.

It was waiting for you to get quiet.

You did.\
"""

CLOSING_C_TITLE = "Closing Reflection C — The God Who Is Near"
CLOSING_C = """\
You will set this book down, and the room will still be the room.

The same light. The same sounds from outside. The same ordinary weight of whatever the day holds next.

And God will be there.

Not waiting for a more spiritual version of you to arrive. Not requiring that you have understood everything. Simply present — the way he has always been present — within the ordinary hours, patient with the pace of ordinary lives.

He met people at wells when they were just trying to get water. He sat at tables with people who were surprised to be invited. He walked beside two people on a road and let them talk for hours before they recognized him.

He has always been patient with that kind of pace.

The room you return to is not a lesser place than the one this book has inhabited. The ordinary moment — the kitchen, the drive, the quiet morning before anyone else is awake — is not a distance from what this book has been pointing toward.

It is exactly where the pointing has been aimed.

You do not need to go looking for him.

He is already present — within whatever ordinary thing you carry into the next hour, and the one after that.

The invitation has not closed.\
"""

CLOSING_MEETING_PLACE_TITLE = "Closing Reflection — The Meeting Place"
CLOSING_MEETING_PLACE = """\
The room is still the room.

The same light. The same sounds. The same ordinary weight of whatever comes next.

Something has shifted in you, briefly. Not a transformation — a recognition.

The book described a meeting. What it could not tell you — what only this moment can — is that you are already in it.

God was not waiting for you to finish reading. The speaking was not something that would begin once you set the book down and became very still. It was already reaching toward you. Through the ordinary things. Through the hours. Through a life that has always been spoken to, even in the years when you weren't attending.

The woman at the well was not looking for an encounter. She was just trying to get water. The two on the road didn't know they were in a meeting until it was nearly over.

He has always preferred that kind of company.

This is not a descent, returning to the ordinary after this. The ordinary room is where the meeting has always been held — not as symbol, not as pathway, but as the meeting place itself.

You do not need to become more ready than you already are.

He was not waiting for you to find him.
He has been walking with you.\
"""

# ─── Document Factory ─────────────────────────────────────────────────────────

def new_doc():
    """Return a fresh Document with manuscript margins and base font."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

    # Default Normal style — Times New Roman 12
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    return doc


def set_font(run, size=12, bold=False, italic=False):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic


def add_empty(doc, size_pt=6):
    """Add a thin spacer paragraph."""
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(size_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing      = Pt(size_pt)
    return p


def add_title_block(doc, title, subtitle=None):
    """Centered title + optional subtitle, followed by a rule-space."""
    doc.add_paragraph()   # top breathing room

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, size=16, bold=True)
    p.paragraph_format.space_after = Pt(4)

    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = s.add_run(subtitle)
        set_font(r2, size=12, italic=True)
        s.paragraph_format.space_after = Pt(0)

    add_empty(doc, 18)


def add_chapter_heading(doc, text):
    """Chapter heading — centered, 14pt bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=14, bold=True)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(24)


def add_section_heading(doc, part_label, heading_text):
    """Part label line + section title, centered."""
    add_empty(doc, 12)

    # Part label (e.g., "Part II")
    pl = doc.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pl.add_run(part_label)
    set_font(r1, size=11, italic=True)
    pl.paragraph_format.space_after = Pt(4)

    # Heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p.add_run(heading_text)
    set_font(r2, size=14, bold=True)
    p.paragraph_format.space_after = Pt(24)


def word_count(text):
    return len(re.findall(r'\S+', text))


def add_prose_poetry(doc, raw_text):
    """
    Parse prose-poetic text:
      - Blank lines  → paragraph breaks (Space After 12pt)
      - Lines within a block → soft line-breaks within one paragraph
      - Lines beginning '- ' → bullet list items
    """
    blocks = re.split(r'\n{2,}', raw_text.strip())

    for block in blocks:
        lines = block.split('\n')

        # Detect whether any line in this block is a bullet item
        has_bullets = any(l.startswith('- ') for l in lines)

        if has_bullets:
            # First, emit non-bullet lines as a single paragraph
            pre_lines = [l for l in lines if not l.startswith('- ')]
            if pre_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(0)
                first = True
                for line in pre_lines:
                    if not first:
                        p.add_run().add_break(WD_BREAK.LINE)
                    r = p.add_run(line)
                    set_font(r)
                    first = False

            # Emit bullet lines
            bullet_lines = [l[2:] for l in lines if l.startswith('- ')]
            for item in bullet_lines:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent  = Inches(0.4)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(2)
                r = p.add_run(item)
                set_font(r)

            add_empty(doc, 10)

        else:
            # Regular prose-poetry block
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(12)
            first = True
            for line in lines:
                if not first:
                    p.add_run().add_break(WD_BREAK.LINE)
                r = p.add_run(line)
                set_font(r)
                first = False


# ─── Individual Document Builders ─────────────────────────────────────────────

def build_chapter18(path_stem):
    doc = new_doc()
    add_title_block(doc, BOOK_TITLE, BOOK_SUBTITLE)
    add_chapter_heading(doc, CHAPTER_18_HEADING)
    add_prose_poetry(doc, CHAPTER_18)
    out = OUTPUT_DIR / (path_stem + ".docx")
    doc.save(out)
    return out


def build_closing(path_stem, title, text):
    doc = new_doc()
    add_title_block(doc, BOOK_TITLE, BOOK_SUBTITLE)
    add_chapter_heading(doc, title)
    add_prose_poetry(doc, text)
    out = OUTPUT_DIR / (path_stem + ".docx")
    doc.save(out)
    return out


def build_combined(path_stem):
    doc = new_doc()

    # ── Cover / title page ──────────────────────────────────────────────────
    add_title_block(doc, BOOK_TITLE, BOOK_SUBTITLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Review Packet")
    set_font(r, size=13, italic=True)
    p.paragraph_format.space_after = Pt(0)

    # ── Part I — Chapter 18 ─────────────────────────────────────────────────
    doc.add_page_break()
    add_section_heading(doc, "Part I", CHAPTER_18_HEADING)
    add_prose_poetry(doc, CHAPTER_18)

    # ── Part II — Closing Reflection A ─────────────────────────────────────
    doc.add_page_break()
    add_section_heading(doc, "Part II", CLOSING_A_TITLE)
    add_prose_poetry(doc, CLOSING_A)

    # ── Part III — Closing Reflection B ────────────────────────────────────
    doc.add_page_break()
    add_section_heading(doc, "Part III", CLOSING_B_TITLE)
    add_prose_poetry(doc, CLOSING_B)

    # ── Part IV — Closing Reflection C ─────────────────────────────────────
    doc.add_page_break()
    add_section_heading(doc, "Part IV", CLOSING_C_TITLE)
    add_prose_poetry(doc, CLOSING_C)

    # ── Part V — The Meeting Place ──────────────────────────────────────────
    doc.add_page_break()
    add_section_heading(doc, "Part V", CLOSING_MEETING_PLACE_TITLE)
    add_prose_poetry(doc, CLOSING_MEETING_PLACE)

    out = OUTPUT_DIR / (path_stem + ".docx")
    doc.save(out)
    return out


# ─── PDF Conversion ───────────────────────────────────────────────────────────

def convert_to_pdf(docx_path):
    """Convert a .docx to .pdf via docx2pdf (uses Word COM on Windows)."""
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path, None
    except Exception as e:
        return None, str(e)


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    tasks = [
        ("chapter_18_condensed_review",              build_chapter18,
         None, None),
        ("closing_reflection_A_laborers_return",     build_closing,
         CLOSING_A_TITLE, CLOSING_A),
        ("closing_reflection_B_tuning_fork",         build_closing,
         CLOSING_B_TITLE, CLOSING_B),
        ("closing_reflection_C_god_who_is_near",     build_closing,
         CLOSING_C_TITLE, CLOSING_C),
        ("closing_reflection_the_meeting_place",         build_closing,
         CLOSING_MEETING_PLACE_TITLE, CLOSING_MEETING_PLACE),
        ("chapter18_and_closing_options_review_packet", build_combined,
         None, None),
    ]

    word_counts = {
        "Chapter 18 (condensed)":          word_count(CHAPTER_18),
        "Closing Reflection A":            word_count(CLOSING_A),
        "Closing Reflection B":            word_count(CLOSING_B),
        "Closing Reflection C":            word_count(CLOSING_C),
        "Closing Reflection — The Meeting Place": word_count(CLOSING_MEETING_PLACE),
    }

    results   = []
    pdf_warns = []

    for stem, builder, title, text in tasks:
        # Build docx
        if title is None:
            docx_path = builder(stem)
        else:
            docx_path = builder(stem, title, text)

        # Convert to pdf
        pdf_path, err = convert_to_pdf(docx_path)
        if err:
            pdf_warns.append(f"  {stem}.pdf — {err}")
            pdf_path = None

        results.append((stem, docx_path, pdf_path))
        print(f"[OK] {stem}.docx" + (" + .pdf" if pdf_path else " (PDF failed)"))

    # ── Report ───────────────────────────────────────────────────────────────
    print("\n" + "═" * 60)
    print("GENERATION REPORT")
    print("═" * 60)
    print(f"\nOutput directory:\n  {OUTPUT_DIR}\n")

    print("FILES CREATED:")
    for stem, docx_path, pdf_path in results:
        size_kb = round(docx_path.stat().st_size / 1024, 1)
        print(f"  {docx_path.name}  ({size_kb} KB)")
        if pdf_path and pdf_path.exists():
            size_kb2 = round(pdf_path.stat().st_size / 1024, 1)
            print(f"  {pdf_path.name}  ({size_kb2} KB)")

    print("\nWORD COUNTS (body text only):")
    for label, count in word_counts.items():
        print(f"  {label}: {count} words")
    combined_wc = (word_count(CHAPTER_18) + word_count(CLOSING_A)
                   + word_count(CLOSING_B) + word_count(CLOSING_C)
                   + word_count(CLOSING_MEETING_PLACE))
    print(f"  Combined packet (all body text): {combined_wc} words")

    if pdf_warns:
        print("\nPDF WARNINGS:")
        for w in pdf_warns:
            print(w)
    else:
        print("\nNo PDF warnings.")

    print("\nDone.")


if __name__ == "__main__":
    main()
