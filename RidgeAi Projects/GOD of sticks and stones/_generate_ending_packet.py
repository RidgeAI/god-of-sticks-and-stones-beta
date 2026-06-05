"""
Ending Review Packet Generator
The God of Sticks and Stones — Chapter 18, Closing Reflection, A Final Word
Source: review_edition_literary_refined.md (approved manuscript)
"""

import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK

OUTPUT_DIR = Path(
    r"C:\Users\ridge\GitHub_Projects\RidgeAi Projects\GOD of sticks and stones"
    r"\master_manuscript\07_publication_build\review_packets"
)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Section labels and titles ─────────────────────────────────────────────────

CHAPTER_18_LABEL = "Chapter Eighteen"
CHAPTER_18_TITLE = "The God of Sticks and Stones"

CLOSING_LABEL = "Closing Reflection"
CLOSING_TITLE = "The Laborer’s Return"

FINAL_WORD_LABEL = "A Final Word"
FINAL_WORD_TITLE = "The Meeting Place"

# ── Content — sourced verbatim from review_edition_literary_refined.md ────────

CHAPTER_18 = """\
In the beginning,
humanity stood barefoot upon living ground.

Dust and breath.
Garden and river.
Tree and fruit.

The story began close to the earth.

Before towers.
Before empires.
Before systems and monuments.

There was communion.

And throughout scripture,
God continually speaks through ordinary things.

Stick.
Stone.
Seed.
Field.
Water.
Tree.

Simple things.
Foundational things.
Things humanity often overlooks while searching for greatness elsewhere.

Yet the biblical story continually returns to these humble materials because they reveal the structure of reality itself.

The seed descends before rising.
Roots deepen before branches spread.
Water flows downward.
The shepherd walks low places.
The well is dug beneath the surface.
The cornerstone rests beneath the structure.

Again and again,
creation preaches the same sermon.

Downward before upward.
Humility before exaltation.
Communion before dominion.

Human civilization repeatedly struggles against this pattern.

Cain wanders.
Babel rises.
Empires expand.
Humanity attempts to secure meaning,
identity,
and permanence through the works of its own hands.

Brick replaces stone.
Tower replaces garden.
Control replaces trust.

Yet even then,
God continues speaking.

Through wilderness.
Through shepherds.
Through caves.
Through wells.
Through altars of uncut stone.

The kingdom continually moves through ordinary life rather than spectacle.

Abraham listens beneath trees.
Jacob sleeps upon stone.
Moses carries a staff.
David gathers stones from a river.
Prophets retreat into wilderness.

Then Christ arrives.

Not as emperor.
Not as architect of empire.

A carpenter.

Hands touching wood.
Feet walking dusty roads.
Teaching beside fields,
rivers,
and mountains.

The Creator enters creation without distancing Himself from it.

And eventually,
wood becomes cross.
Stone becomes tomb.
Garden becomes resurrection.

The entire biblical story converges there.

The tree of grasping in Eden becomes the tree of surrender at Calvary.
The stone sealing death becomes the stone rolled away.
The lowered place becomes the place where life rises again.

Creation itself testifies.

The earth shakes.
Rocks split.
Darkness covers the land.

The stones preach.

Perhaps this is why Christ said that if humanity remained silent,
the stones themselves would cry out.

Creation has always witnessed the story.

The stars preach scale and wonder.
Fields preach patience.
Water preaches dependence.
Gravity preaches humility.
Trees preach rootedness.
Seeds preach resurrection.

Everything speaks.

Modern civilization often becomes too loud to hear these things clearly.

Humanity surrounds itself with systems,
noise,
speed,
and endless distraction.

People increasingly live within environments almost entirely shaped by human hands.
Concrete.
Glass.
Screens.
Artificial light.
Digital identities.

The older language becomes difficult to hear.

Yet beneath every system,
reality remains unchanged.

Human beings still thirst.
Still hunger.
Still grieve.
Still long for meaning.
Still require love.
Still depend upon breath.

No civilization has outgrown dependence.

This is why the kingdom continually calls humanity back toward humility.

Not humiliation.
Not self-hatred.

Truthful rootedness.

The recognition that life itself is received.

Modern humanity often seeks identity through visibility,
performance,
accumulation,
and external validation.

But scripture repeatedly reveals that the deepest life develops underground first.

Roots remain hidden.
Wells remain buried.
Seeds disappear into darkness before emergence.

The kingdom grows quietly.

This quietness often appears weak to civilizations organized around spectacle.

But creation itself reveals that hidden foundations sustain visible life.

A tree without roots falls.
A tower without foundation cracks.
A field without water dies.

The deeper realities always matter most.

This is why scripture continually warns about pride.

Pride is not merely arrogance.

Pride is disconnection from reality.

The illusion that humanity is self-created,
self-sustaining,
and ultimately self-governing.

Yet gravity eventually humbles every tower.

Kings die.
Empires collapse.
Structures decay.
Bodies return to dust.

The ground receives everyone equally.

And still,
scripture never approaches humanity with contempt.

This matters profoundly.

The field is not hated.
The harvest is not despised.

Human beings remain deeply loved despite wandering.

The laborers are sent not to condemn the world,
but to cultivate within it.

Christ eats with sinners.
Touches the sick.
Weeps over cities.
Offers living water to the thirsty.

The kingdom moves toward brokenness rather than away from it.

Because the deepest problem was never merely ignorance.

It was separation.

Humanity became disconnected:
from God,
from creation,
from one another,
and even from itself.

Modern civilization often intensifies this fragmentation.

People become isolated while constantly connected.
Informed while inwardly exhausted.
Visible while spiritually unseen.

The soul becomes crowded with noise.

Yet beneath the noise,
creation continues speaking patiently.

The stars still move overhead.
Rivers still carve through stone.
Seeds still enter the earth.
Rain still falls upon fields.

Reality itself still points toward communion.

Perhaps this is why the biblical story ends not with escape from creation,
but with restoration within it.

A river flows again.
The tree of life returns.
The city descends.
Communion is restored.

The final vision unites everything humanity once divided.

Garden and city.
Stone and life.
Cultivation and worship.
Creation and civilization.

The false separation heals.

This reveals something extraordinary.

God never abandoned creation.

The world was never merely disposable scenery.

The sticks and stones mattered.
The fields mattered.
The rivers mattered.
The trees mattered.

Not because they were divine in themselves,
but because they continually revealed the character and order of the One who made them.

The farmer working the field,
the shepherd guiding sheep,
the laborer digging a well,
the carpenter shaping wood—
all became living parables of the kingdom.

Perhaps modern humanity has not truly become more advanced than previous generations.

Perhaps humanity has simply become more distracted.

The pace increased.
The towers grew taller.
The systems became more complex.

But the deepest human questions remain unchanged.

What is life?
What is truth?
What is worth building?
What foundation can endure?
What truly satisfies thirst?

Scripture answers these questions not first through abstraction,
but through living symbols grounded in creation itself.

Water.
Seed.
Bread.
Stone.
Cross.
River.
Tree.

Again and again,
God speaks through things close enough to touch.

This keeps the kingdom near the ground.

Near enough for children.
Near enough for laborers.
Near enough for the poor.
Near enough for farmers,
shepherds,
and ordinary people.

Empire often speaks the language of complexity.
The kingdom often speaks the language of creation.

Not because truth is simplistic,
but because reality itself is deeply integrated.

The stars and the seed tell the same story.
The well and the cross tell the same story.
The river and the tree tell the same story.

Life comes through communion,
through surrender,
through rootedness,
and through receiving what humanity cannot manufacture for itself.

This is why Christ becomes the cornerstone.

Not merely another teacher,
but reality rightly aligned.

The rejected stone.
The living water.
The true vine.
The good shepherd.
The bread of life.

All the symbols converge in Him.
In Jesus.

And still,
He remains close enough to kneel beside the ground.

Perhaps this is the final invitation of the kingdom.

Not upward striving toward self-made transcendence.

But downward rootedness into reality,
communion,
and love.

To become once again:
- living soil,
- faithful laborers,
- rooted trees,
- people willing to receive living water,
- builders who remember the cornerstone,
- and stewards humble enough to walk with God through the field.

For beneath every empire,
beyond every tower,
and underneath all the noise of human history,
the older realities still remain.

Stone.
Tree.
Water.
Seed.
Dust.
Breath.

And the God who has always spoken through them.\
"""

CLOSING_A = """\
The work is still waiting for you.

Not the extraordinary kind — just the ordinary. The meetings and the dishes. The conversation that doesn’t quite resolve. The commute. The quiet morning. The neighbor whose name you keep meaning to learn.

You have been walking through something slow and old. And now the ordinary day is in front of you again, with its small weight, asking for your presence.

This is not a descent.

Those who worked close to the ground understood that faithfulness rarely meant stepping away from ordinary life. It meant bringing a different quality of attention into it. Showing up. Tending what was given to them to tend.

You don’t have to be extraordinary at this.

You only need to remain present.

And the one who worked with his hands, who ate with people who weren’t certain they belonged at the table, has always been willing to meet you in the ordinary hours — in the commute, in the kitchen, in the work that will still be there tomorrow morning.

He did not require spectacle from the people he met along the road.

He is not requiring it now.

So go back to the work.

Quietly, and with purpose — the way a laborer returns after a long rest. The task was always ordinary. And you were always enough to carry it.\
"""

CLOSING_MEETING_PLACE = """\
The room is still the room.

The same light. The same sounds. The same ordinary weight of whatever comes next.

Something has shifted in you, briefly. Not a transformation — a recognition.

The book described a meeting. What it could not tell you — what only this moment can — is that you are already in it.

God was not waiting for you to finish reading. The speaking was not something that would begin once you set the book down and became very still. It was already reaching toward you. Through the ordinary things. Through the hours. Through a life that has always been spoken to, even in the years when you weren’t attending.

The woman at the well was not looking for an encounter. She was just trying to get water. The two on the road didn’t know they were in a meeting until it was nearly over.

He has always preferred that kind of company.

This is not a descent, returning to the ordinary after this. The ordinary room is where the meeting has always been held — not as symbol, not as pathway, but as the meeting place itself.

You do not need to become more ready than you already are.

He was not waiting for you to find him.
He has been walking with you.\
"""

# ── Document helpers ──────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return doc


def set_font(run, size=12, bold=False, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic


def add_empty(doc, size_pt=6):
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(size_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing      = Pt(size_pt)
    return p


def add_section_heading(doc, label, title, top_space=True):
    if top_space:
        add_empty(doc, 12)

    pl = doc.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pl.add_run(label)
    set_font(r1, size=11, italic=True)
    pl.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p.add_run(title)
    set_font(r2, size=14, bold=True)
    p.paragraph_format.space_after = Pt(24)


def add_prose_poetry(doc, raw_text):
    blocks = re.split(r'\n{2,}', raw_text.strip())

    for block in blocks:
        lines = block.split('\n')
        has_bullets = any(l.startswith('- ') for l in lines)

        if has_bullets:
            pre_lines = [l for l in lines if not l.startswith('- ')]
            if pre_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(4)
                p.paragraph_format.space_before = Pt(0)
                first = True
                for line in pre_lines:
                    if not first:
                        p.add_run().add_break(WD_BREAK.LINE)
                    r = p.add_run(line)
                    set_font(r)
                    first = False

            for item in [l[2:] for l in lines if l.startswith('- ')]:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent  = Inches(0.4)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(2)
                r = p.add_run(item)
                set_font(r)

            add_empty(doc, 10)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(12)
            first = True
            for line in lines:
                if not first:
                    p.add_run().add_break(WD_BREAK.LINE)
                r = p.add_run(line)
                set_font(r)
                first = False


# ── Packet builder ────────────────────────────────────────────────────────────

def build_packet(stem):
    doc = new_doc()

    # Chapter 18
    add_section_heading(doc, CHAPTER_18_LABEL, CHAPTER_18_TITLE, top_space=False)
    add_prose_poetry(doc, CHAPTER_18)

    # Closing Reflection
    doc.add_page_break()
    add_section_heading(doc, CLOSING_LABEL, CLOSING_TITLE, top_space=True)
    add_prose_poetry(doc, CLOSING_A)

    # A Final Word
    doc.add_page_break()
    add_section_heading(doc, FINAL_WORD_LABEL, FINAL_WORD_TITLE, top_space=True)
    add_prose_poetry(doc, CLOSING_MEETING_PLACE)

    out = OUTPUT_DIR / (stem + ".docx")
    doc.save(out)
    return out


def convert_to_pdf(docx_path):
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path, None
    except Exception as e:
        return None, str(e)


def get_pdf_page_count(pdf_path):
    try:
        import PyPDF2
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return len(reader.pages)
    except Exception:
        pass
    try:
        from pypdf import PdfReader
        with open(pdf_path, "rb") as f:
            reader = PdfReader(f)
            return len(reader.pages)
    except Exception:
        pass
    return None


def get_docx_page_count(docx_path):
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        pages = doc.ComputeStatistics(2)  # wdStatisticPages
        doc.Close(False)
        word.Quit()
        return pages
    except Exception:
        return None


def fmt_ts(path):
    return datetime.fromtimestamp(path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    stem = "chapter18_to_end_review_packet"
    print(f"Building {stem}.docx ...")
    docx_path = build_packet(stem)
    print(f"  [OK] DOCX saved")

    print(f"Converting to PDF ...")
    pdf_path, err = convert_to_pdf(docx_path)
    if err:
        print(f"  [WARN] PDF conversion error: {err}")
    else:
        print(f"  [OK] PDF saved")

    print()
    print("=" * 70)
    print("VERIFICATION REPORT")
    print("=" * 70)
    print()

    # DOCX
    docx_pages = get_docx_page_count(docx_path)
    print("DOCX")
    print(f"  Path:      {docx_path}")
    print(f"  Size:      {docx_path.stat().st_size:,} bytes")
    print(f"  Modified:  {fmt_ts(docx_path)}")
    print(f"  Pages:     {docx_pages if docx_pages else 'requires Word rendering'}")
    print()

    # PDF
    if pdf_path and pdf_path.exists():
        pdf_pages = get_pdf_page_count(pdf_path)
        print("PDF")
        print(f"  Path:      {pdf_path}")
        print(f"  Size:      {pdf_path.stat().st_size:,} bytes")
        print(f"  Modified:  {fmt_ts(pdf_path)}")
        print(f"  Pages:     {pdf_pages if pdf_pages else 'unavailable'}")
        print()

    print("CONTENTS")
    print("  [1] Chapter 18 — The God of Sticks and Stones")
    print("      Includes: All the symbols converge in Him. / In Jesus.")
    print("  [2] Closing Reflection — The Laborer's Return")
    print("  [3] A Final Word — The Meeting Place")
    print("      Ends: He was not waiting for you to find him.")
    print("            He has been walking with you.")
    print()
    print("  No other content.")
    print()
    print("Done.")


if __name__ == "__main__":
    main()
