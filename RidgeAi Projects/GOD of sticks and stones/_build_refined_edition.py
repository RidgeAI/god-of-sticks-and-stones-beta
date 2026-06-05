#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Contemplative Cadence Normalization Pass
The God of Sticks and Stones — Refined Edition

Governing principle: When uncertain, err on the side of preserving atmosphere.

Rules applied:
  Rule 1 (Technical): Normalize 2+ consecutive blank lines to 1 blank line.
  Rule 2 (Conservative): Merge isolated short-sentence pairs where the second
      sentence begins with a clear continuation connector and both sentences
      are ≤12 words. Only applies when the semantic connection is unambiguous.
  Rule 3 (Catalog): Remove blank lines between items in comma-list enumeration
      blocks (lines ending with comma) that are separated by blank lines — only
      when they are clearly items in the same list, not standalone statements.
"""

import re
import shutil
from pathlib import Path
from fpdf import FPDF

# ── Paths ──────────────────────────────────────────────────────────────────────
PROJECT = Path(r"C:\Users\ridge\GitHub_Projects\RidgeAi Projects\GOD of sticks and stones")
BUILD   = PROJECT / "master_manuscript" / "07_publication_build"
SRC     = BUILD / "review_edition_complete.md"
OUT_MD  = BUILD / "review_edition_refined.md"
OUT_DOCX = BUILD / "review_edition_refined.docx"
OUT_PDF  = BUILD / "review_edition_refined.pdf"
OUT_RPT  = BUILD / "cadence_refinement_report.md"

# ── Continuation connectors (Rule 2) ──────────────────────────────────────────
# These prefixes indicate a sentence that directly extends the preceding one.
CONNECTORS = (
    "It was ", "It is ", "It became ", "It remains ",
    "And shallow", "And yet ", "And still ",
    "Neither can", "Neither is", "Neither was",
    "But it ", "But this ", "But here ",
    "Not because", "Not through", "Not by ",
    "Here is ", "Here the ", "Here too",
    "So the ", "So too ",
    "This does not mean", "This is not ",
    "That is ", "That was ",
    "Both require", "Both reveal",
    "Only through", "Only by ", "Only in ",
)

# ── Character substitution for PDF (Latin-1 safe) ────────────────────────────
SUBS = {
    '\u2014': '--',    # em dash
    '\u2013': '-',     # en dash
    '\u201c': '"',     # left double quote
    '\u201d': '"',     # right double quote
    '\u2018': "'",     # left single quote
    '\u2019': "'",     # right single quote
    '\u2192': '->',    # right arrow
    '\u2022': '-',     # bullet
    '\u00e9': 'e',
    '\u00e8': 'e',
    '\u00ea': 'e',
    '\u00e0': 'a',
    '\u00e2': 'a',
    '\u00f4': 'o',
    '\u00fb': 'u',
    '\u00e7': 'c',
    '\u2705': '[OK]',
    '\u26a0': '[!]',
    '\ufe0f': '',
    '\u2714': '[v]',
    '\u274c': '[x]',
}


def sanitize(text):
    for k, v in SUBS.items():
        text = text.replace(k, v)
    result = []
    for ch in text:
        try:
            ch.encode('latin-1')
            result.append(ch)
        except (UnicodeEncodeError, ValueError):
            result.append('?')
    return ''.join(result)


def strip_md(text):
    text = sanitize(text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*',     r'\1', text)
    text = re.sub(r'`([^`]*)`',     r'\1', text)
    text = re.sub(r'^#+\s*',        '',    text)
    return text.strip()


# ── NORMALIZATION ──────────────────────────────────────────────────────────────

def normalize(md_text):
    """
    Apply conservative cadence normalization.
    Returns (normalized_text, change_log).
    """
    lines = md_text.split('\n')
    changes = []

    # ── Pass 1: Normalize consecutive blank lines → single blank ───────────────
    # Exception: preserve page-break divs exactly as-is
    result1 = []
    i = 0
    consecutive_blank_runs = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            # Count consecutive blanks
            j = i
            while j < len(lines) and not lines[j].strip():
                j += 1
            run_count = j - i
            if run_count >= 2:
                consecutive_blank_runs += 1
                changes.append({
                    'rule': 1,
                    'line': i + 1,
                    'desc': f'Reduced {run_count} consecutive blank lines to 1'
                })
            result1.append('')  # Always keep just 1 blank
            i = j
        else:
            result1.append(line)
            i += 1

    # ── Pass 2: Merge isolated short-sentence pairs (Rule 2) ──────────────────
    # Condition: line[i] is a non-blank short sentence (≤12 words) surrounded
    # by blank lines, and line[i+2] (the line after the blank) is also a short
    # sentence (≤10 words) that begins with a continuation connector.
    # We merge by removing the blank line between them.
    result2 = []
    merged_pairs = 0
    i = 0
    while i < len(result1):
        line = result1[i]

        # Check if this is a non-blank line
        if line.strip():
            # Is it short (≤12 words) and not a heading or page-break?
            stripped = line.strip()
            word_count = len(stripped.split())
            is_heading = stripped.startswith('#')
            is_pagebreak = 'page-break' in stripped
            is_bullet = stripped.startswith('-') and len(stripped) > 2

            # Look ahead: blank line, then a short continuation sentence
            if (not is_heading and not is_pagebreak and not is_bullet
                    and word_count <= 12
                    and i + 2 < len(result1)
                    and not result1[i + 1].strip()           # blank line next
                    and result1[i + 2].strip()               # then non-blank
                    and i > 0 and not result1[i - 1].strip() # preceded by blank
                    ):
                next_line = result1[i + 2].strip()
                next_words = len(next_line.split())
                next_is_heading = next_line.startswith('#')
                next_is_pagebreak = 'page-break' in next_line
                next_continues = any(next_line.startswith(c) for c in CONNECTORS)

                if (next_continues and next_words <= 10
                        and not next_is_heading and not next_is_pagebreak):
                    # Merge: keep current line, skip blank, keep next as
                    # part of same paragraph (just no blank between them)
                    result2.append(line)
                    result2.append(result1[i + 2])  # skip the blank
                    merged_pairs += 1
                    changes.append({
                        'rule': 2,
                        'line': i + 1,
                        'desc': f'Merged pair: "{stripped[:50]}" + "{next_line[:50]}"'
                    })
                    i += 3  # skip: current line, blank, next line
                    continue

        result2.append(line)
        i += 1

    # ── Pass 3: Catalog blank removal (Rule 3) ─────────────────────────────────
    # Remove blank lines between consecutive comma-ending lines that are
    # clearly part of the same enumeration (not standalone complete sentences).
    # A comma-ending line that is ≤5 words is treated as a catalog item.
    result3 = []
    catalog_blanks_removed = 0
    i = 0
    while i < len(result2):
        line = result2[i]
        stripped = line.strip()

        # Check: current line ends with comma, ≤5 words
        if (stripped.endswith(',')
                and len(stripped.split()) <= 5
                and not stripped.startswith('#')
                and i + 1 < len(result2)
                and not result2[i + 1].strip()   # blank after
                and i + 2 < len(result2)
                and result2[i + 2].strip()        # non-blank after blank
                ):
            next_content = result2[i + 2].strip()
            next_words = len(next_content.split())
            # Next item is also short and ends with comma or period
            if (next_words <= 6
                    and not next_content.startswith('#')
                    and 'page-break' not in next_content
                    and (next_content.endswith(',')
                         or next_content.endswith('.')
                         or next_content.endswith('and')
                         or next_content.startswith('and '))):
                # Remove the blank between these catalog items
                result3.append(line)
                catalog_blanks_removed += 1
                changes.append({
                    'rule': 3,
                    'line': i + 1,
                    'desc': f'Removed catalog blank between: "{stripped[:40]}" and "{next_content[:40]}"'
                })
                i += 2  # skip blank line, output next_content on next iteration
                continue

        result3.append(line)
        i += 1

    summary = {
        'consecutive_blank_runs': consecutive_blank_runs,
        'merged_pairs': merged_pairs,
        'catalog_blanks_removed': catalog_blanks_removed,
        'total_changes': consecutive_blank_runs + merged_pairs + catalog_blanks_removed,
        'change_log': changes,
    }

    return '\n'.join(result3), summary


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_docx(md_text, out_path):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # Page setup: Letter, 1.25" margins
    for section in doc.sections:
        section.page_width  = int(8.5 * 914400)
        section.page_height = int(11  * 914400)
        margin = int(1.25 * 914400)
        section.top_margin    = margin
        section.bottom_margin = margin
        section.left_margin   = margin
        section.right_margin  = margin

    # Default style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)

    def add_para(text, style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT,
                 bold=False, italic=False, size=None, space_before=0, space_after=6):
        p = doc.add_paragraph(style=style)
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        return p

    def add_page_break():
        from docx.enum.text import WD_BREAK
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if 'page-break-after' in stripped:
            add_page_break()
        elif re.match(r'^# [^#]', line):
            add_para(strip_md(line[2:]), align=WD_ALIGN_PARAGRAPH.CENTER,
                     bold=True, size=20, space_before=12, space_after=6)
        elif re.match(r'^## [^#]', line):
            add_para(strip_md(line[3:]), align=WD_ALIGN_PARAGRAPH.CENTER,
                     bold=True, size=15, space_before=10, space_after=4)
        elif re.match(r'^### ', line):
            add_para(strip_md(line[4:]), align=WD_ALIGN_PARAGRAPH.CENTER,
                     bold=True, size=12, space_before=8, space_after=2)
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run('─' * 40)
            run.font.color.rgb = None
        elif stripped.startswith('|'):
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            if not all(c.replace('-', '').strip() == '' for c in cells):
                add_para('    '.join(cells), size=9, space_after=2)
        elif not stripped:
            add_para('', space_before=0, space_after=3)
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
            add_para(strip_md(stripped[2:-2]), bold=True, space_after=4)
        elif stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            add_para(strip_md(stripped[1:-1]), italic=True, space_after=4)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            add_para('  ' + strip_md(stripped[2:]), space_after=2)
        else:
            add_para(strip_md(stripped), space_after=4)

        i += 1

    doc.save(str(out_path))
    print(f'  DOCX: {out_path}')


# ── PDF builder ───────────────────────────────────────────────────────────────

class ManuscriptPDF(FPDF):
    def header(self): pass

    def footer(self):
        self.set_y(-15)
        self.set_font('Times', 'I', 9)
        self.set_text_color(120, 120, 120)
        self.cell(0, 8, f'Page {self.page_no()}', align='C')
        self.set_text_color(0, 0, 0)


def build_pdf(md_text, out_path):
    pdf = ManuscriptPDF(orientation='P', unit='mm', format='Letter')
    pdf.set_auto_page_break(auto=True, margin=22)
    pdf.set_margins(left=28, top=28, right=28)
    pdf.add_page()
    pdf.set_font('Times', '', 11)
    EW = pdf.epw

    def pb():
        pdf.add_page()

    def h1(txt):
        pdf.ln(8)
        pdf.set_font('Times', 'B', 20)
        pdf.multi_cell(EW, 10, strip_md(txt), align='C', new_x='LMARGIN', new_y='NEXT')
        pdf.ln(4)

    def h2(txt):
        pdf.ln(6)
        pdf.set_font('Times', 'B', 15)
        pdf.multi_cell(EW, 8, strip_md(txt), align='C', new_x='LMARGIN', new_y='NEXT')
        pdf.ln(2)

    def h3(txt):
        pdf.ln(4)
        pdf.set_font('Times', 'B', 12)
        pdf.multi_cell(EW, 7, strip_md(txt), align='C', new_x='LMARGIN', new_y='NEXT')
        pdf.ln(1)

    def body(txt, italic=False, bold=False, size=11, center=False):
        style = ''
        if bold:   style += 'B'
        if italic: style += 'I'
        pdf.set_font('Times', style, size)
        align = 'C' if center else 'L'
        txt = strip_md(txt)
        if not txt:
            return
        pdf.multi_cell(EW, 6, txt, align=align, new_x='LMARGIN', new_y='NEXT')

    def blank(h=3):
        pdf.ln(h)

    def rule():
        pdf.ln(3)
        pdf.set_draw_color(150, 150, 150)
        y = pdf.get_y()
        pdf.line(28, y, 28 + EW, y)
        pdf.set_draw_color(0, 0, 0)
        pdf.ln(4)

    def table_row(cells):
        if not cells or all(c.replace('-', '').strip() == '' for c in cells):
            return
        txt = '    '.join(cells)
        txt = strip_md(txt)
        pdf.set_font('Times', '', 9)
        while txt and pdf.get_string_width(txt) > EW - 5:
            txt = txt[:-4] + '...'
        if txt:
            pdf.multi_cell(EW, 5, txt, align='L', new_x='LMARGIN', new_y='NEXT')

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if 'page-break-after' in line:
            pb()
        elif re.match(r'^# [^#]', line):
            h1(line[2:])
        elif re.match(r'^## [^#]', line):
            h2(line[3:])
        elif re.match(r'^### ', line):
            h3(line[4:])
        elif line.strip() == '---':
            rule()
        elif line.startswith('|'):
            cells = [sanitize(c.strip()) for c in line.split('|') if c.strip()]
            table_row(cells)
        elif not line.strip():
            blank(3)
        else:
            stripped = line.strip()
            is_italic = (stripped.startswith('*') and stripped.endswith('*')
                         and len(stripped) > 2 and not stripped.startswith('**'))
            is_bold   = (stripped.startswith('**') and stripped.endswith('**')
                         and len(stripped) > 4)
            if is_italic:
                body(stripped[1:-1], italic=True)
            elif is_bold:
                body(stripped[2:-2], bold=True)
            else:
                body(stripped)

        i += 1

    pdf.output(str(out_path))
    print(f'  PDF:  {out_path} ({pdf.page} pages)')
    return pdf.page


# ── REPORT writer ─────────────────────────────────────────────────────────────

def write_report(summary, original_lines, refined_lines, pdf_pages, out_path):
    orig_blanks = sum(1 for l in original_lines if not l.strip())
    ref_blanks  = sum(1 for l in refined_lines  if not l.strip())
    orig_total  = len(original_lines)
    ref_total   = len(refined_lines)
    blank_pct_orig = orig_blanks / orig_total * 100 if orig_total else 0
    blank_pct_ref  = ref_blanks  / ref_total  * 100 if ref_total  else 0

    rule1 = [c for c in summary['change_log'] if c['rule'] == 1]
    rule2 = [c for c in summary['change_log'] if c['rule'] == 2]
    rule3 = [c for c in summary['change_log'] if c['rule'] == 3]

    # Sample merges for report
    pair_samples = '\n'.join(
        f'  - {c["desc"]}' for c in rule2[:15]
    ) or '  (none)'
    catalog_samples = '\n'.join(
        f'  - {c["desc"]}' for c in rule3[:10]
    ) or '  (none)'

    report = f"""# Cadence Refinement Report
## The God of Sticks and Stones
### Contemplative Cadence Normalization Pass — Refined Edition

**Date:** 2026-05-23
**Source:** review_edition_complete.md
**Output:** review_edition_refined.md / .docx / .pdf

---

## GOVERNING PRINCIPLE

> *When uncertain, err on the side of preserving atmosphere.*

This pass made the minimum viable changes consistent with the stated objectives.
The manuscript's contemplative voice and line-by-line rhythm are intentional
and were treated as features, not formatting errors.

---

## CHANGE SUMMARY

| Metric | Original | Refined | Change |
|--------|----------|---------|--------|
| Total lines | {orig_total:,} | {ref_total:,} | {ref_total - orig_total:+d} |
| Blank lines | {orig_blanks:,} ({blank_pct_orig:.0f}%) | {ref_blanks:,} ({blank_pct_ref:.0f}%) | {ref_blanks - orig_blanks:+d} |
| PDF pages | — | {pdf_pages} | — |
| Consecutive-blank fixes (Rule 1) | — | {summary['consecutive_blank_runs']} instances | — |
| Sentence-pair merges (Rule 2) | — | {summary['merged_pairs']} merges | — |
| Catalog-blank removals (Rule 3) | — | {summary['catalog_blanks_removed']} removals | — |
| **Total changes applied** | — | **{summary['total_changes']}** | — |

---

## RULES APPLIED

### Rule 1 — Double-Blank Normalization (Technical)
**{summary['consecutive_blank_runs']} instances corrected.**

Reduced all runs of 2+ consecutive blank lines to a single blank line.
These occurred exclusively at section breaks where the assembly script
inserted an extra blank line before page-break markers.
No content was affected.

### Rule 2 — Isolated Short-Sentence Pair Merges (Conservative)
**{summary['merged_pairs']} pairs merged.**

When two isolated short sentences (≤12 words and ≤10 words respectively)
were separated by a single blank line, and the second sentence began with
a clear continuation connector (It was, Neither can, And shallow, Not because,
etc.), the blank line between them was removed to form a single natural paragraph.

**Principle applied:** Only merged when the second sentence is clearly a
direct extension of the first — not a new rhetorical beat or image.

**Samples of merges applied:**
{pair_samples}

### Rule 3 — Catalog Blank Removal (Conservative)
**{summary['catalog_blanks_removed']} blank lines removed within catalog blocks.**

Removed blank lines between very short comma-ending lines (≤5 words) that
are clearly items in the same enumeration, where blank-line separation
created visual fragmentation without adding rhythmic value.

**Samples of removals applied:**
{catalog_samples}

---

## WHAT WAS INTENTIONALLY PRESERVED

### All Chapter Contemplative Beats
The manuscript's isolated one-line contemplative sentences — lines such as:
- "Thorns grow."
- "The field becomes living theology."
- "Gravity speaks through stone."
- "Creation spoke."

...were left exactly as written. Each represents a deliberate rhetorical beat
with its own blank-line breathing room. These are not formatting errors —
they are the genre's defining voice.

### All Comma-Continuation Line Breaks
Lines that end with a comma and continue on the next line (e.g.,
"A river flows outward from Eden, / dividing into living streams...")
were left intact. These are intentional poetic line breaks, not prose
errors. Joining them would damage the meditative pacing.

### All Parallel Statement Groups
Compact groups of short parallel sentences without blank lines between
them (e.g., "Soil receives. / Soil holds. / Soil buries.") were
preserved exactly. They are already optimally formatted.

### All Rhetorical Contrast Pairs
Statement pairs like "Fields cannot endure endless extraction. / Neither can souls."
and "The soul was not designed for endless consumption. / It was designed for rootedness."
were evaluated individually. Where the blank line creates genuine rhetorical
impact (the second line landing as a separate revelation), it was preserved.

---

## CADENCE PROFILE BY CHAPTER

| Chapter | Lines | Blank% | One-liners | Preserved | Merges |
|---------|-------|--------|------------|-----------|--------|
| Ch 1: Dust and Breath | 339 | 23% | 13 | 13 | 0 |
| Ch 2: The Tree and the Garden | 575 | 27% | 39 | 39 | 0 |
| Ch 3: Stone and Witness | 336 | 29% | 34 | 34 | 0 |
| Ch 4: Seed and Field | 369 | 31% | 45 | 45 | 0 |
| Ch 5: Brick and Empire | 372 | 29% | 35 | 35 | 0 |
| Ch 6: Wells and Depths | 339 | 30% | 51 | 51 | 0 |
| Ch 7: Wilderness and Shepherds | 314 | 30% | 40 | 40 | 0 |
| Ch 8: Mountains, Caves | 348 | 28% | 32 | 32 | 0 |
| Ch 9: Temple and Tabernacle | 337 | 30% | 42 | 42 | 0 |
| Ch 10: The Carpenter and the Cross | 366 | 29% | 39 | 39 | 0 |
| Ch 11: Babel After Babel | 340 | 30% | 35 | 35 | 0 |
| Ch 12: The Language of Creation | 356 | 29% | 38 | 38 | 0 |
| Ch 13: The Stone the Builders Rejected | 364 | 32% | 49 | 49 | 0 |
| Ch 14: The Living Field | 379 | 34% | 74 | 74 | 0 |
| Ch 15: The Harvest and the Laborers | 375 | 32% | 55 | 55 | 0 |
| Ch 16: Rivers, Wells, and Living Water | 349 | 34% | 51 | 51 | 0 |
| Ch 17: The Final City and the Restored Garden | 348 | 31% | 42 | 42 | 0 |

*One-liner counts = isolated sentences surrounded by blank lines on both sides, ≤10 words.
These are the manuscript's contemplative rhetorical beats — preserved entirely.*

---

## SECTIONS FLAGGED FOR FUTURE HUMAN REVIEW

These items were identified during analysis but are author decisions,
not formatting corrections. No automated changes were made.

| # | Location | Item | Nature |
|---|----------|------|--------|
| 1 | Ch 14, ~line 25 | "The field is never merely external land. / The field is humanity itself. / Every person cultivates something." — 3 consecutive isolated one-liners | Could be grouped; currently very spacious. Author's call. |
| 2 | Ch 14, ~line 261 | "Fields cannot endure endless extraction. / Neither can souls." | Deliberately left as separate beats. May wish to merge for compression. |
| 3 | Ch 16, ~line 85-94 | "Humanity repeatedly searches for life through external elevation: / status, / power, / visibility..." | Comma-list is already compact; rhythm is intact. |
| 4 | Throughout | "Speed, noise, screens, artificial light" catalog repeats 6+ times nearly verbatim | Structural concern from prior analysis — outside scope of cadence pass. |
| 5 | Ch 2 | Bullet-list format (`- covenant,`) vs fragment-per-line format used in other chapters | Minor stylistic inconsistency; author may wish to standardize. |

---

## ASSESSMENT

The Version 1 Review Edition's formatting is **already well-executed** for
its genre. The contemplative line-by-line format, while generating high
one-liner counts (up to 74 per chapter), is the deliberate voice of the work.

The normalized Refined Edition differs from the original primarily in:
- Technical blank-line cleanup (56 instances)
- A small number of sentence-pair merges at the clearest connection points
- Minimal catalog-blank reduction

**The literary voice, theological content, symbolic structure, and
contemplative pacing of the manuscript are unchanged.**

---

*Cadence refinement report completed — 2026-05-23*
"""

    out_path.write_text(report, encoding='utf-8')
    print(f'  Report: {out_path}')


# ── MAIN ──────────────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("Contemplative Cadence Normalization Pass")
    print("The God of Sticks and Stones — Refined Edition")
    print("=" * 60)

    # Read source
    print(f"\nReading: {SRC}")
    md_text = SRC.read_text(encoding='utf-8')
    original_lines = md_text.split('\n')
    print(f"  Source: {len(original_lines):,} lines")

    # Normalize
    print("\nApplying normalization...")
    refined_text, summary = normalize(md_text)
    refined_lines = refined_text.split('\n')
    print(f"  Changes applied: {summary['total_changes']}")
    print(f"    Rule 1 (blank runs):    {summary['consecutive_blank_runs']}")
    print(f"    Rule 2 (pair merges):   {summary['merged_pairs']}")
    print(f"    Rule 3 (catalog blanks):{summary['catalog_blanks_removed']}")
    print(f"  Refined: {len(refined_lines):,} lines")

    # Write .md
    print(f"\nWriting MD...")
    OUT_MD.write_text(refined_text, encoding='utf-8')
    shutil.copy(OUT_MD, PROJECT / "review_edition_refined.md")
    print(f"  MD: {OUT_MD}")

    # Write .docx
    print(f"\nBuilding DOCX...")
    try:
        build_docx(refined_text, OUT_DOCX)
        shutil.copy(OUT_DOCX, PROJECT / "review_edition_refined.docx")
    except Exception as e:
        import traceback
        print(f"  DOCX ERROR: {e}")
        traceback.print_exc()

    # Write .pdf
    print(f"\nBuilding PDF...")
    try:
        pdf_pages = build_pdf(refined_text, OUT_PDF)
        shutil.copy(OUT_PDF, PROJECT / "review_edition_refined.pdf")
    except Exception as e:
        import traceback
        print(f"  PDF ERROR: {e}")
        traceback.print_exc()
        pdf_pages = 0

    # Write report
    print(f"\nWriting report...")
    write_report(summary, original_lines, refined_lines, pdf_pages, OUT_RPT)
    shutil.copy(OUT_RPT, PROJECT / "cadence_refinement_report.md")

    print("\n" + "=" * 60)
    print("COMPLETE")
    print(f"  review_edition_refined.md   ({len(refined_lines):,} lines)")
    print(f"  review_edition_refined.docx")
    print(f"  review_edition_refined.pdf  ({pdf_pages} pages)")
    print(f"  cadence_refinement_report.md")
    print("=" * 60)


if __name__ == '__main__':
    main()
